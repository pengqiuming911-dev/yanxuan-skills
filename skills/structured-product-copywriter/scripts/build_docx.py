#!/usr/bin/env python3
"""按固定模板把推介材料组装成 Word(.docx)。

输入一个 manifest JSON,描述文档各段(标题/正文/参数块/图片/分隔/文案文件)。
build_docx.py 据此渲染 Word。图片路径找不到时插入"[图片待补:xxx]"占位,不报错。

manifest 示例:
{
  "title": "中证1000 2倍DCN 推介材料",
  "sections": [
    {"type":"heading","text":"产品结构文字版（长版）"},
    {"type":"copy_file","path":"copy.txt"},
    {"type":"heading","text":"产品派息与敲出观察点位表"},
    {"type":"image","path":"product-card.png","caption":"产品结构解析卡"},
    {"type":"heading","text":"产品胜率数据"},
    {"type":"body","text":"回测区间 2016-06-26 至 2026-06-25,胜率 98.14%"},
    {"type":"image","path":"winrate.png","caption":"通毓终端回测结果"}
  ]
}

section 类型:
  heading     一级加粗标题
  subheading  二级标题
  body        正文段落(按 \\n 分段)
  params      参数块(等宽小字,按行)
  image       插图(等比缩放到页宽,缺文件则占位)
  separator   分隔线
  copy_file   读文案文件,按行自动判别 标题/参数/正文/分隔
"""
import argparse
import json
import os
import sys
import tempfile

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_chinese_font(doc, font_name="Microsoft YaHei"):
    """给 Normal 样式设中文字体(含 eastAsia),保证中文正常显示。"""
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15 if level == 1 else 12.5)
    p.space_before = Pt(8)
    p.space_after = Pt(4)
    return p


def add_body(doc, text, size=11):
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.size = Pt(size)


def add_params(doc, text, size=10):
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.size = Pt(size)
            r.font.name = "Consolas"  # 参数块用等宽,对齐整齐
            rpr = r._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            from docx.oxml import OxmlElement
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")


# 大图嵌入前自动压缩,避免 docx 体积撑爆飞书 upload-docx 的 nginx body 限制
# (AMAC 全页公示图常 1MB+、一万像素高,原样嵌入 → 上传 413)。阈值、目标宽高、JPEG 质量见下。
COMPRESS_THRESHOLD = 400 * 1024  # 超过 400KB 才压
COMPRESS_MAX_WIDTH = 1100        # 宽上限 1100px(16cm 印刷够清晰)
COMPRESS_MAX_HEIGHT = 4000       # 高上限 4000px:AMAC 全页图一万像素高,不限高压完还是大
COMPRESS_JPEG_QUALITY = 82


def _maybe_compress(full):
    """大图先压成 JPEG(宽≤1100、高≤4000, q82)再返回,原文件不动。

    返回 (embed_path, is_temp):is_temp=True 时 embed_path 是临时文件,调用方用完要删。
    打不开/不识别就原样返回,别阻断装配流程——宁可图大也别让整个 docx 生成失败。
    """
    try:
        size = os.path.getsize(full)
    except OSError:
        return full, False
    if size <= COMPRESS_THRESHOLD:
        return full, False
    from PIL import Image
    try:
        im = Image.open(full).convert("RGB")
    except Exception:
        return full, False
    w, h = im.size
    # 等比缩到宽、高都不超上限(以更紧的那个为准),AMAC 全页图靠高上限兜住
    scale = min(COMPRESS_MAX_WIDTH / w, COMPRESS_MAX_HEIGHT / h, 1.0)
    if scale < 1.0:
        im = im.resize((max(1, int(w * scale)), max(1, int(h * scale))), Image.LANCZOS)
    fd, tmp = tempfile.mkstemp(suffix=".jpg")
    os.close(fd)
    im.save(tmp, "JPEG", quality=COMPRESS_JPEG_QUALITY, optimize=True)
    return tmp, True


def add_image(doc, path, caption=None, width_cm=16.0, max_height_cm=20.0, base_dir=None):
    full = path
    if base_dir and not os.path.isabs(full):
        full = os.path.join(base_dir, path)
    if not os.path.exists(full):
        p = doc.add_paragraph(f"[图片待补: {path}]")
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xC0, 0x50, 0x50)
        return
    # 大图先压缩(见 _maybe_compress),避免 docx 体积超飞书上传 body 限制
    embed_path, is_temp = _maybe_compress(full)
    # 按宽 16cm 等比缩放;若过高(超 max_height_cm)则改为按高度缩放,避免撑爆多页
    from PIL import Image
    with Image.open(embed_path) as im:
        iw, ih = im.size
    w = width_cm
    h = w * ih / iw
    if h > max_height_cm:
        h = max_height_cm
        w = h * iw / ih
    doc.add_picture(embed_path, width=Cm(w), height=Cm(h))
    if is_temp:
        try:
            os.remove(embed_path)
        except OSError:
            pass
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def add_separator(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "CCCCCC")
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_hyperlink(paragraph, url, text):
    """在段落里加一个可点击的超链接(蓝色下划线)。"""
    from docx.oxml import OxmlElement
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rPr.append(rFonts)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_link_list(doc, items, size=11):
    """items: [{label, url}]。每项一段:label + ": " + 超链接(url)。url 为空则只显示 label。"""
    for it in items:
        label = it.get("label", "")
        url = it.get("url", "")
        p = doc.add_paragraph()
        if label:
            r = p.add_run(label + ":")
            r.font.size = Pt(size)
        if url:
            if label:
                p.add_run(" ")
            add_hyperlink(p, url, url)
        for r in p.runs:
            r.font.size = Pt(size)


def render_copy_file(doc, path, base_dir=None):
    full = path
    if base_dir and not os.path.isabs(full):
        full = os.path.join(base_dir, path)
    if not os.path.exists(full):
        add_body(doc, f"[文案文件缺失: {path}]")
        return
    with open(full, encoding="utf-8") as f:
        lines = [l.rstrip("\n") for l in f]
    # 去掉 BMP 外字符(emoji 等),Word 某些字体不含
    lines = ["".join(ch for ch in l if ord(ch) <= 0xFFFF) for l in lines]
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s == "---":
            add_separator(doc)
            continue
        if s.startswith("🚀") or (s.startswith("【") and s.endswith("】")):
            add_heading(doc, s, level=2)
        elif s.endswith("：") or ("：" in s[:6]):
            add_params(doc, s)
        else:
            add_body(doc, s)


def build(manifest_path, output_path):
    with open(manifest_path, encoding="utf-8") as f:
        manifest = json.load(f)
    base_dir = os.path.dirname(os.path.abspath(manifest_path))

    doc = Document()
    set_chinese_font(doc)
    # 页边距
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)

    if manifest.get("title"):
        p = doc.add_paragraph(manifest["title"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(18)
        doc.add_paragraph()

    for sec in manifest.get("sections", []):
        t = sec.get("type")
        if t == "heading":
            add_heading(doc, sec["text"], level=1)
        elif t == "subheading":
            add_heading(doc, sec["text"], level=2)
        elif t == "body":
            add_body(doc, sec["text"], size=sec.get("size", 11))
        elif t == "params":
            add_params(doc, sec["text"])
        elif t == "image":
            add_image(doc, sec["path"], sec.get("caption"), base_dir=base_dir)
        elif t == "separator":
            add_separator(doc)
        elif t == "link_list":
            add_link_list(doc, sec.get("items", []))
        elif t == "copy_file":
            render_copy_file(doc, sec["path"], base_dir=base_dir)
        else:
            add_body(doc, str(sec.get("text", "")))

    doc.save(output_path)
    print(f"Word 生成: {output_path}")


def copy_to_desktop(src_path):
    """把生成的 Word 复制一份到桌面,方便用户找。Windows 桌面 = ~/Desktop。"""
    import shutil
    home = os.path.expanduser("~")
    desktop = os.path.join(home, "Desktop")
    if not os.path.isdir(desktop):
        # 某些中文 Windows 桌面目录名不同,试 OneDrive 桌面
        alt = os.path.join(home, "OneDrive", "Desktop")
        if os.path.isdir(alt):
            desktop = alt
        else:
            print(f"找不到桌面目录(试过 {desktop}),跳过复制。")
            return None
    dst = os.path.join(desktop, os.path.basename(src_path))
    try:
        # 桌面已有同名文件且被 Word 打开会锁,先尝试,失败就换个名
        shutil.copy2(src_path, dst)
        print(f"已复制到桌面: {dst}")
        return dst
    except PermissionError:
        base, ext = os.path.splitext(dst)
        i = 2
        while True:
            cand = f"{base}-{i}{ext}"
            try:
                shutil.copy2(src_path, cand)
                print(f"桌面同名文件被占用,已另存: {cand}")
                return cand
            except PermissionError:
                i += 1
                if i > 20:
                    print("桌面同名文件多次占用,跳过。")
                    return None


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--manifest", required=True, help="manifest JSON 路径")
    ap.add_argument("--output", required=True, help="输出 .docx 路径")
    ap.add_argument("--desktop", action="store_true", help="生成后复制一份到桌面")
    args = ap.parse_args()
    build(args.manifest, args.output)
    if args.desktop:
        copy_to_desktop(args.output)


if __name__ == "__main__":
    main()
